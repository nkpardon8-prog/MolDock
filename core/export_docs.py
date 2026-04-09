"""
MoleCopilot document export utilities.

Converts structured data and Markdown text into DOCX, PDF, and XLSX
reports.  All functions return the absolute path of the written file
as a plain ``str`` so callers can hand it to downstream tools or embed
it in JSON responses.
"""

from __future__ import annotations

import datetime
import re
from pathlib import Path
from typing import Optional

from core.utils import setup_logging, REPORTS_DIR, RESULTS_DIR, ensure_dir, load_env

logger = setup_logging("export_docs")

# ---------------------------------------------------------------------------
# Internal Markdown parsing helpers
# ---------------------------------------------------------------------------

_RE_HEADING = re.compile(r"^(#{1,6})\s+(.*)")
_RE_BOLD = re.compile(r"\*\*(.+?)\*\*")
_RE_TABLE_ROW = re.compile(r"^\|(.+)\|$")
_RE_TABLE_SEP = re.compile(r"^\|[\s\-:]+\|$")
_RE_LIST_ITEM = re.compile(r"^\s*[-*]\s+(.*)")


def _default_output_path(output_path: Optional[str], suffix: str, title: Optional[str]) -> Path:
    """Derive a deterministic output path inside ``REPORTS_DIR``.

    Parameters
    ----------
    output_path : str or None
        Explicit caller-provided path.  When ``None`` a timestamped name
        is generated.
    suffix : str
        File extension including the dot (e.g. ``".docx"``).
    title : str or None
        Optional title used to build a human-friendly filename.

    Returns
    -------
    Path
        Resolved output path.
    """
    if output_path is not None:
        dest = Path(output_path)
        ensure_dir(dest.parent)
        return dest
    ensure_dir(REPORTS_DIR)
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    stem = re.sub(r"[^\w]+", "_", (title or "report").strip().lower())[:60]
    return REPORTS_DIR / f"{stem}_{ts}{suffix}"


def _parse_markdown_blocks(markdown_text: str) -> list[dict]:
    """Split Markdown text into a list of typed blocks.

    Each block is a dict with ``"type"`` (``"heading"``, ``"table"``,
    ``"list"``, ``"paragraph"``) and type-specific payload keys.

    Parameters
    ----------
    markdown_text : str
        Raw Markdown source.

    Returns
    -------
    list[dict]
        Ordered sequence of blocks.
    """
    blocks: list[dict] = []
    lines = markdown_text.split("\n")
    idx = 0

    while idx < len(lines):
        line = lines[idx]

        # Empty line — skip
        if not line.strip():
            idx += 1
            continue

        # Heading
        heading_match = _RE_HEADING.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            blocks.append({"type": "heading", "level": level, "text": text})
            idx += 1
            continue

        # Table — collect consecutive rows
        if _RE_TABLE_ROW.match(line.strip()):
            rows: list[list[str]] = []
            while idx < len(lines) and _RE_TABLE_ROW.match(lines[idx].strip()):
                row_line = lines[idx].strip()
                if _RE_TABLE_SEP.match(row_line):
                    idx += 1
                    continue
                cells = [c.strip() for c in row_line.strip("|").split("|")]
                rows.append(cells)
                idx += 1
            blocks.append({"type": "table", "rows": rows})
            continue

        # List item(s)
        list_match = _RE_LIST_ITEM.match(line)
        if list_match:
            items: list[str] = []
            while idx < len(lines) and _RE_LIST_ITEM.match(lines[idx]):
                m = _RE_LIST_ITEM.match(lines[idx])
                if m:
                    items.append(m.group(1))
                idx += 1
            blocks.append({"type": "list", "items": items})
            continue

        # Plain paragraph — collect until blank or structural line
        para_lines: list[str] = []
        while idx < len(lines):
            cur = lines[idx]
            if not cur.strip():
                break
            if _RE_HEADING.match(cur) or _RE_TABLE_ROW.match(cur.strip()) or _RE_LIST_ITEM.match(cur):
                break
            para_lines.append(cur)
            idx += 1
        blocks.append({"type": "paragraph", "text": " ".join(para_lines)})

    return blocks


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------


def export_docx(
    markdown_text: str,
    output_path: Optional[str] = None,
    figures: Optional[list[str]] = None,
    title: Optional[str] = None,
) -> str:
    """Export Markdown text to a ``.docx`` Word document.

    Supports headings (``#`` through ``######``), **bold** runs, bullet
    lists, pipe-delimited tables, and optional figure insertion.  When
    *title* is supplied a simple title page is prepended.

    Parameters
    ----------
    markdown_text : str
        Source Markdown.
    output_path : str, optional
        Explicit destination path.  Defaults to a timestamped file in
        ``REPORTS_DIR``.
    figures : list[str], optional
        Paths to image files to append at the end of the document.
    title : str, optional
        Document title shown on the first page.

    Returns
    -------
    str
        Absolute path of the saved ``.docx`` file.
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # -- Title page -----------------------------------------------------------
    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(26)

        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = subtitle_para.add_run(
            f"Generated by MoleCopilot — {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        date_run.font.size = Pt(11)
        date_run.italic = True

        doc.add_page_break()

    # -- Render Markdown blocks -----------------------------------------------
    blocks = _parse_markdown_blocks(markdown_text)

    for block in blocks:
        btype = block["type"]

        if btype == "heading":
            level: int = min(block["level"], 4)  # python-docx supports 0-9
            doc.add_heading(block["text"], level=level)

        elif btype == "paragraph":
            _add_rich_paragraph(doc, block["text"])

        elif btype == "list":
            for item in block["items"]:
                _add_rich_paragraph(doc, item, style="List Bullet")

        elif btype == "table":
            rows: list[list[str]] = block["rows"]
            if not rows:
                continue
            num_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = "Light Grid Accent 1"
            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    if ci < num_cols:
                        table.rows[ri].cells[ci].text = cell_text.strip()
            doc.add_paragraph()  # spacing after table

    # -- Figures --------------------------------------------------------------
    if figures:
        doc.add_page_break()
        doc.add_heading("Figures", level=2)
        for fig_path_str in figures:
            fig_path = Path(fig_path_str)
            if fig_path.is_file():
                doc.add_picture(str(fig_path), width=Inches(5.5))
                cap = doc.add_paragraph(fig_path.stem)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
            else:
                logger.warning("Figure not found, skipping: %s", fig_path_str)

    # -- Save -----------------------------------------------------------------
    dest = _default_output_path(output_path, ".docx", title)
    doc.save(str(dest))
    logger.info("DOCX saved: %s", dest)
    return str(dest.resolve())


def _add_rich_paragraph(doc: object, text: str, style: Optional[str] = None) -> None:
    """Add a paragraph to *doc*, rendering ``**bold**`` spans as bold runs.

    Parameters
    ----------
    doc : Document
        The python-docx ``Document`` instance.
    text : str
        Plain text that may contain ``**bold**`` markers.
    style : str, optional
        Named paragraph style (e.g. ``"List Bullet"``).
    """
    para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = _RE_BOLD.split(text)
    # split yields: [normal, bold_content, normal, bold_content, ...]
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        if i % 2 == 1:
            run.bold = True


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------


def export_pdf(
    markdown_text: str,
    output_path: Optional[str] = None,
    figures: Optional[list[str]] = None,
    title: Optional[str] = None,
) -> str:
    """Export Markdown text to a PDF document using fpdf2.

    Supports headings, bold text, bullet lists, tables, and figure
    insertion.

    Parameters
    ----------
    markdown_text : str
        Source Markdown.
    output_path : str, optional
        Explicit destination path.  Defaults to a timestamped file in
        ``REPORTS_DIR``.
    figures : list[str], optional
        Paths to image files to embed in the PDF.
    title : str, optional
        Document title shown on the first page.

    Returns
    -------
    str
        Absolute path of the saved ``.pdf`` file.
    """
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Register a Unicode-capable built-in font for body text
    pdf.set_font("Helvetica", size=11)

    # -- Title page -----------------------------------------------------------
    if title:
        pdf.set_font("Helvetica", "B", 24)
        pdf.cell(0, 40, title, new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.set_font("Helvetica", "I", 11)
        pdf.cell(
            0, 10,
            f"Generated by MoleCopilot  --  {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}",
            new_x="LMARGIN", new_y="NEXT",
            align="C",
        )
        pdf.add_page()
        pdf.set_font("Helvetica", size=11)

    # -- Render blocks --------------------------------------------------------
    blocks = _parse_markdown_blocks(markdown_text)

    for block in blocks:
        btype = block["type"]

        if btype == "heading":
            level: int = block["level"]
            font_size = max(18 - (level - 1) * 2, 11)
            pdf.set_font("Helvetica", "B", font_size)
            pdf.ln(4)
            pdf.multi_cell(0, font_size * 0.6, block["text"])
            pdf.ln(2)
            pdf.set_font("Helvetica", size=11)

        elif btype == "paragraph":
            _pdf_rich_text(pdf, block["text"])
            pdf.ln(3)

        elif btype == "list":
            for item in block["items"]:
                pdf.cell(8)  # indent
                pdf.cell(5, 6, "-")  # bullet marker
                _pdf_rich_text(pdf, item, indent=13)
                pdf.ln(2)
            pdf.ln(2)

        elif btype == "table":
            _pdf_table(pdf, block["rows"])
            pdf.ln(4)

    # -- Figures --------------------------------------------------------------
    if figures:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, "Figures", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", size=11)
        page_w = pdf.w - pdf.l_margin - pdf.r_margin
        for fig_path_str in figures:
            fig_path = Path(fig_path_str)
            if fig_path.is_file():
                # Fit image within page width
                pdf.image(str(fig_path), w=min(page_w, 170))
                pdf.set_font("Helvetica", "I", 9)
                pdf.cell(0, 6, fig_path.stem, new_x="LMARGIN", new_y="NEXT", align="C")
                pdf.set_font("Helvetica", size=11)
                pdf.ln(4)
            else:
                logger.warning("Figure not found, skipping: %s", fig_path_str)

    # -- Save -----------------------------------------------------------------
    dest = _default_output_path(output_path, ".pdf", title)
    pdf.output(str(dest))
    logger.info("PDF saved: %s", dest)
    return str(dest.resolve())


def _pdf_rich_text(pdf: object, text: str, indent: float = 0) -> None:
    """Write *text* to the PDF, rendering ``**bold**`` spans.

    Parameters
    ----------
    pdf : FPDF
        The active FPDF instance.
    text : str
        Text that may contain ``**bold**`` markers.
    indent : float
        Horizontal offset from the left margin (points).
    """
    parts = _RE_BOLD.split(text)
    if indent:
        pdf.cell(indent)
    for i, part in enumerate(parts):
        if not part:
            continue
        if i % 2 == 1:
            pdf.set_font("Helvetica", "B", 11)
        else:
            pdf.set_font("Helvetica", size=11)
        # Use write() so text wraps naturally
        pdf.write(6, part)
    pdf.ln()


def _pdf_table(pdf: object, rows: list[list[str]]) -> None:
    """Render a table into the PDF as a simple cell grid.

    Parameters
    ----------
    pdf : FPDF
        The active FPDF instance.
    rows : list[list[str]]
        Table rows; the first row is treated as a header.
    """
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    if num_cols == 0:
        return

    page_w = pdf.w - pdf.l_margin - pdf.r_margin
    col_w = page_w / num_cols

    for ri, row in enumerate(rows):
        for ci in range(num_cols):
            cell_text = row[ci].strip() if ci < len(row) else ""
            if ri == 0:
                pdf.set_font("Helvetica", "B", 10)
            else:
                pdf.set_font("Helvetica", size=10)
            pdf.cell(col_w, 7, cell_text, border=1)
        pdf.ln()
    pdf.set_font("Helvetica", size=11)


# ---------------------------------------------------------------------------
# XLSX export
# ---------------------------------------------------------------------------


def export_xlsx(
    data: dict,
    output_path: Optional[str] = None,
    title: Optional[str] = None,
) -> str:
    """Export a dict of record-lists to a multi-sheet Excel workbook.

    Each key in *data* becomes a worksheet name and the corresponding
    value (a ``list[dict]``) is written as a table with auto-sized
    columns.  If any column name contains the substring ``"energy"``
    (case-insensitive) a red-green conditional colour scale is applied
    so that more negative (better) binding energies stand out.

    Parameters
    ----------
    data : dict
        Mapping of ``sheet_name -> list[dict]``.
    output_path : str, optional
        Explicit destination.  Defaults to a timestamped file in
        ``REPORTS_DIR``.
    title : str, optional
        Human-readable title used only for the filename stem.

    Returns
    -------
    str
        Absolute path of the saved ``.xlsx`` file.
    """
    import pandas as pd
    from openpyxl import load_workbook
    from openpyxl.formatting.rule import ColorScaleRule
    from openpyxl.utils import get_column_letter

    dest = _default_output_path(output_path, ".xlsx", title)

    # -- Write each sheet with pandas -----------------------------------------
    with pd.ExcelWriter(str(dest), engine="openpyxl") as writer:
        for sheet_name, records in data.items():
            if not records:
                # Write an empty sheet with a note
                df = pd.DataFrame({"Note": ["No data available"]})
            else:
                df = pd.DataFrame(records)
            safe_name = re.sub(r"[\\/*?\[\]:]", "_", sheet_name)[:31]
            df.to_excel(writer, sheet_name=safe_name, index=False)

    # -- Post-process: auto-width + conditional formatting --------------------
    wb = load_workbook(str(dest))

    for ws in wb.worksheets:
        # Auto-fit column widths
        for col_cells in ws.columns:
            max_length = 0
            col_letter = get_column_letter(col_cells[0].column)
            for cell in col_cells:
                try:
                    cell_len = len(str(cell.value)) if cell.value is not None else 0
                    if cell_len > max_length:
                        max_length = cell_len
                except Exception:
                    pass
            adjusted = min(max_length + 4, 50)
            ws.column_dimensions[col_letter].width = adjusted

        # Conditional formatting for "energy" columns
        header_row = [cell.value for cell in ws[1]]
        for ci, header in enumerate(header_row, start=1):
            if header and "energy" in str(header).lower():
                col_letter = get_column_letter(ci)
                data_range = f"{col_letter}2:{col_letter}{ws.max_row}"
                rule = ColorScaleRule(
                    start_type="min",
                    start_color="63BE7B",   # green (best / most negative)
                    mid_type="percentile",
                    mid_value=50,
                    mid_color="FFEB84",     # yellow
                    end_type="max",
                    end_color="F8696B",     # red (worst / least negative)
                )
                ws.conditional_formatting.add(data_range, rule)
                logger.info(
                    "Added energy colour-scale to column %s ('%s') in sheet '%s'",
                    col_letter, header, ws.title,
                )

    wb.save(str(dest))
    logger.info("XLSX saved: %s", dest)
    return str(dest.resolve())


# ---------------------------------------------------------------------------
# Standalone demo
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import json

    print("=" * 70)
    print("MoleCopilot Export Module — Standalone Demo")
    print("=" * 70)

    sample_markdown = """\
# MoleCopilot Docking Report

## Target Information

The target protein is **Aromatase** (CYP19A1), a key enzyme in estrogen
biosynthesis.  It is a validated drug target for **breast cancer** therapy.

## Docking Results

| Rank | Compound | Binding Energy (kcal/mol) | Ki (nM) |
|------|----------|---------------------------|---------|
| 1 | Letrozole | -9.8 | 12.3 |
| 2 | Anastrozole | -9.2 | 24.1 |
| 3 | Exemestane | -8.7 | 45.6 |
| 4 | Fadrozole | -8.3 | 67.2 |

## Key Observations

- **Letrozole** shows the strongest predicted binding affinity
- The triazole ring forms a coordination bond with the heme iron
- All top compounds occupy the active-site cavity near the I-helix
- Hydrophobic contacts with Phe221 and Val370 are conserved across poses

## Conclusion

Virtual screening confirms known aromatase inhibitors as top-ranked
hits.  The binding poses are consistent with published crystal structures.
"""

    # -- DOCX demo ------------------------------------------------------------
    print("\n--- DOCX Export ---")
    try:
        docx_path = export_docx(
            markdown_text=sample_markdown,
            title="MoleCopilot Docking Report",
        )
        print(f"  Created: {docx_path}")
    except Exception as exc:
        logger.error("DOCX export failed: %s", exc)
        print(f"  [SKIPPED] DOCX export failed: {exc}")

    # -- PDF demo -------------------------------------------------------------
    print("\n--- PDF Export ---")
    try:
        pdf_path = export_pdf(
            markdown_text=sample_markdown,
            title="MoleCopilot Docking Report",
        )
        print(f"  Created: {pdf_path}")
    except Exception as exc:
        logger.error("PDF export failed: %s", exc)
        print(f"  [SKIPPED] PDF export failed: {exc}")

    # -- XLSX demo ------------------------------------------------------------
    print("\n--- XLSX Export ---")
    try:
        xlsx_data = {
            "Docking Results": [
                {"Rank": 1, "Compound": "Letrozole", "Binding Energy": -9.8, "Ki (nM)": 12.3},
                {"Rank": 2, "Compound": "Anastrozole", "Binding Energy": -9.2, "Ki (nM)": 24.1},
                {"Rank": 3, "Compound": "Exemestane", "Binding Energy": -8.7, "Ki (nM)": 45.6},
                {"Rank": 4, "Compound": "Fadrozole", "Binding Energy": -8.3, "Ki (nM)": 67.2},
            ],
            "ADMET Properties": [
                {"Compound": "Letrozole", "LogP": 2.5, "MW": 285.3, "HBA": 4, "HBD": 0, "TPSA": 78.3},
                {"Compound": "Anastrozole", "LogP": 1.8, "MW": 293.4, "HBA": 4, "HBD": 0, "TPSA": 78.3},
                {"Compound": "Exemestane", "LogP": 3.1, "MW": 296.4, "HBA": 2, "HBD": 0, "TPSA": 34.1},
            ],
        }
        xlsx_path = export_xlsx(data=xlsx_data, title="Docking Analysis")
        print(f"  Created: {xlsx_path}")
    except Exception as exc:
        logger.error("XLSX export failed: %s", exc)
        print(f"  [SKIPPED] XLSX export failed: {exc}")

    print("\n" + "=" * 70)
    print("Demo complete.")
